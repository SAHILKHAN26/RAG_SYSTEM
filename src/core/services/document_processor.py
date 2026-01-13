"""
Document processing module for extracting text from various file formats.

Supports:
- PDF files (using PyPDF2 and pdfplumber)
- Text files (.txt, .md)
- Word documents (.docx)
"""

import os
from typing import Optional, Tuple
from pathlib import Path

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# DOCX processing
try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

from src.core.logger import app_logger as logger
from src.core.constant import ALLOWED_FILE_TYPES, MAX_FILE_SIZE


class DocumentProcessor:
    """Process and extract text from various document formats"""
    
    def __init__(self):
        """Initialize the document processor"""
        self.supported_types = ALLOWED_FILE_TYPES
        logger.info(f"Document processor initialized. Supported types: {self.supported_types}")
    
    def validate_file(self, file_path: str, file_size: int) -> Tuple[bool, Optional[str]]:
        """
        Validate file type and size.
        
        Args:
            file_path: Path to the file
            file_size: Size of the file in bytes
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file size
        if file_size > MAX_FILE_SIZE:
            return False, f"File size ({file_size} bytes) exceeds maximum allowed size ({MAX_FILE_SIZE} bytes)"
        
        # Check file extension
        file_ext = Path(file_path).suffix.lower().lstrip('.')
        if file_ext not in self.supported_types:
            return False, f"File type '.{file_ext}' not supported. Allowed types: {self.supported_types}"
        
        # Check if file exists
        if not os.path.exists(file_path):
            return False, f"File not found: {file_path}"
        
        return True, None
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text from a document.
        
        Args:
            file_path: Path to the document
            file_type: Type of the document (pdf, txt, docx, md)
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file type is not supported
            Exception: If extraction fails
        """
        file_type = file_type.lower().lstrip('.')
        
        try:
            if file_type == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type in ['txt', 'md']:
                return self._extract_from_text(file_path)
            elif file_type == 'docx':
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}", exc_info=True)
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        if not PDF_SUPPORT:
            raise ImportError("PDF support not available. Install PyPDF2 and pdfplumber.")
        
        text_parts = []
        
        # Try with pdfplumber first (better for complex PDFs)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            if text_parts:
                logger.info(f"Extracted {len(text_parts)} pages from PDF using pdfplumber")
                return "\n\n".join(text_parts)
        
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed, trying PyPDF2: {e}")
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            logger.info(f"Extracted {len(text_parts)} pages from PDF using PyPDF2")
            return "\n\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {e}", exc_info=True)
            raise
    
    def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    logger.info(f"Extracted text file using {encoding} encoding")
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
        
        except Exception as e:
            logger.error(f"Error reading text file: {e}", exc_info=True)
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not DOCX_SUPPORT:
            raise ImportError("DOCX support not available. Install python-docx.")
        
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            logger.info(f"Extracted {len(text_parts)} paragraphs/rows from DOCX")
            return "\n\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"Error reading DOCX file: {e}", exc_info=True)
            raise
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove excessive newlines
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def process_document(self, file_path: str, file_type: str) -> Tuple[str, int]:
        """
        Process a document: extract and clean text.
        
        Args:
            file_path: Path to the document
            file_type: Type of the document
            
        Returns:
            Tuple of (cleaned_text, character_count)
        """
        # Extract text
        raw_text = self.extract_text(file_path, file_type)
        
        # Clean text
        cleaned_text = self.clean_text(raw_text)
        
        char_count = len(cleaned_text)
        logger.info(f"Processed document: {char_count} characters")
        
        return cleaned_text, char_count


# Global document processor instance
_document_processor_instance: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get or create the global document processor instance"""
    global _document_processor_instance
    if _document_processor_instance is None:
        _document_processor_instance = DocumentProcessor()
    return _document_processor_instance
